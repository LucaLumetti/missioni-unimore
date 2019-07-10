from itertools import zip_longest

import io
import os
import re
from PyPDF2 import PdfFileReader, PdfFileWriter
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseBadRequest
from django.shortcuts import redirect
from docx import Document
from docx.shared import Cm
from reportlab.pdfgen import canvas

from Rimborsi import settings
from .forms import *
from .views import load_json


@login_required
def genera_pdf(request, id):
    if request.method == 'POST':
        moduli_missione = ModuliMissione.objects.get(missione_id=id)
        moduli_missione_form = ModuliMissioneForm(request.POST, instance=moduli_missione)
        if moduli_missione_form.is_valid():
            moduli_missione_form.save()
            # return redirect('profile')

        compila_parte_1(request, id)
        compila_parte_2(request, id)
        compila_autorizz_dottorandi(request, id)
        return redirect('resoconto', id)
    else:
        return HttpResponseBadRequest()


def compila_parte_1(request, id):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, settings.STATIC_URL[1:], 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'ModuloMissione_Part1.pdf')
    coords_dict = {
        "struttura_appartenenza": [70, 748],
        "cognome": [170, 695],
        "nome": [370, 695],
        "luogo_nascita": [100, 681],
        "data_nascita": [400, 681],
        "cf": [133, 654],
        "domicilio": [150, 630],
        "prov": [481, 630],
        "qualifica": [145, 604],
        "datore": [347, 604],
        "destinazione": [165, 578],
        "motivazione": [165, 552],
        "inizio_ora": [140, 513],
        "inizio": [310, 513],
        "durata": [260, 486],
        "fondo": [218, 470],
        "mezzo": [133, 438],
        "motivazione_mezzi": [59, 408],
        "targa": [391, 360],
        "data_richiesta": [133, 180],
    }
    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)
    trasporto = Trasporto.objects.filter(missione=missione)

    trasporto_set = set()
    for t in trasporto:
        trasporto_set.add(t.mezzo)

    targa = ""
    motivazione_auto = ""
    if "AUTO" in eval(missione.mezzi_previsti):
        targa = missione.automobile.targa
        motivazione_auto = ' + '.join(t.lower() for t in eval(missione.motivazione_automobile))

    value_dict = {
        "struttura_appartenenza": missione.struttura_fondi,
        "cognome": profile.user.last_name,
        "nome": profile.user.first_name,
        "luogo_nascita": profile.luogo_nascita.name,
        "data_nascita": profile.data_nascita.strftime('%d/%m/%Y'),
        "cf": profile.cf,
        "domicilio": profile.domicilio_fiscale,
        "prov": profile.domicilio_fiscale_provincia,
        "qualifica": profile.qualifica,
        "datore": profile.datore_lavoro,
        "destinazione": missione.citta_destinazione + ", " + missione.stato_destinazione.nome,
        "motivazione": missione.motivazione,
        "inizio_ora": missione.inizio_ora.strftime('%H:%M'),
        "inizio": missione.inizio.strftime('%d/%m/%Y'),
        "durata": str(missione.durata_gg.days),
        "fondo": missione.fondo,
        # "mezzo": ' + '.join(t for t in trasporto_set),
        "mezzo": ' + '.join(t for t in eval(missione.mezzi_previsti)),
        "motivazione_mezzi": motivazione_auto,
        "targa": targa,
        "data_richiesta": date_richiesta.parte_1.strftime('%d/%m/%Y'),
    }

    buffer = io.BytesIO()
    can = canvas.Canvas(buffer)
    can.setFont("Times-Roman", 11)
    # Write values
    for k, v in coords_dict.items():
        if k == 'cf':
            can.setFont('Courier', 18.2)
            can.drawString(*v, value_dict[k])
            can.setFont("Times-Roman", 11)
        else:
            if value_dict[k] is None:
                value_dict[k] = ""
            can.drawString(*v, value_dict[k])

    can.showPage()
    can.save()
    buffer.seek(0)
    new_pdf = PdfFileReader(buffer)

    # Leggo il file base
    input = PdfFileReader(open(input_file, "rb"))  # Base file
    page = input.getPage(0)
    # Faccio il merge delle modifiche con il file base
    page.mergePage(new_pdf.getPage(0))

    # Scrivo tutto in un file temporaneo
    output = PdfFileWriter()
    output.addPage(page)
    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_parte_1_tmp.pdf')
    outputStream = open(output_name_tmp, "wb")
    output.write(outputStream)
    outputStream.close()

    # Salvo il pdf appena creato dentro a un FileField
    output_name = f'Missione_{missione.id}_parte_1.pdf'
    outputStream = open(output_name_tmp, "rb")
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    moduli_missione.parte_1_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def compila_parte_2(request, id):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, settings.STATIC_URL[1:], 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'ModuloMissione_Part2.docx')
    document = Document(input_file)
    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)
    trasporto = Trasporto.objects.filter(missione=missione)

    trasporto_set = set()
    for t in trasporto:
        trasporto_set.add(t.mezzo)

    targa = ""
    if "AUTO" in trasporto_set:
        targa = missione.automobile.targa

    class ParConfig:
        def __init__(self):
            self.config = []
            self.counter = 0

        def append(self, index, values, excludes=[]):
            self.config.append([index, values, excludes])

        def __getitem__(self, index):
            return self.config[index]

    config = ParConfig()
    config.append('Il sottoscritto', [f'{profile.user.first_name} {profile.user.last_name}'])
    config.append('DICHIARA di aver compiuto la missione a', [
        f'{missione.citta_destinazione} - {missione.stato_destinazione.nome}',
        missione.inizio_ora.strftime('%H:%M'),
        missione.inizio.strftime('%d/%m/%Y'),
        missione.fine_ora.strftime('%H:%M'),
        missione.fine.strftime('%d/%m/%Y')])
    config.append('DICHIARA di aver ricevuto', ['aaa'])
    config.append('nel caso di utilizzo di mezzo proprio', ['aaa'])  # todo mancano km in model),
    config.append('che il costo del biglietto', ['cosa ci va?'])  # todo cosa ci va?),
    config.append('che l’originale del fattura/ricevuta cumulativa', ['aaa', 'aaa', 'aaa'])
    config.append('che il costo della fattura/ricevuta _', ['aa'])
    config.append('che l’originale del fattura/ricevuta cumulativa, relativo a ___', ['aaa', 'aaa', 'aaa'])
    config.append('che l’originale del fattura/ricevuta cumulativa, relativo a __________________________ ',
                  ['aaa', 'aaa', 'aaa'])
    config.append('che il costo della fattura/ricevuta __________________________ ', ['aa'])
    config.append('Data richiesta', [date_richiesta.parte_2.strftime('%d/%m/%Y')])

    for k, values, excludes in config:
        str = ''
        for par in document.paragraphs:
            if k in par.text:
                for s1, s2 in zip_longest(re.sub('_+', '_', par.text).split('_'), values, fillvalue=''):
                    if s2 != '':
                        str += f'{s1}__{s2}__'
                    else:
                        str += f'{s1}'
                for r in par.runs:
                    if len(r.text) > 0:
                        r.text = ''
                par.add_run(text=str)
                break

    # Tabella `VIAGGIO E TRASPORTO`
    table = document.tables[0]

    while len(table.rows) <= len(trasporto):
        row = table.add_row()
        # row.height = Cm(0.61)

    for i, t in enumerate(trasporto, start=1):
        table.cell(i, 0).text = t.data.strftime('%d/%m/%Y')
        table.cell(i, 1).text = f'da {t.da}'
        table.cell(i, 2).text = f'da {t.a}'
        table.cell(i, 3).text = t.mezzo
        table.cell(i, 4).text = t.tipo_costo
        table.cell(i, 5).text = f'{t.costo:.2f}'
        table.rows[i].height = Cm(0.61)

    db_dict = {
        'pernottamento': [],
        'scontrino': [],  # pasti
        'convegno': [],
        'altrespese': [],
    }

    # Load the default values for each field in db_dict
    for k, _ in db_dict.items():
        db_dict[k] = load_json(missione, k)

    r_scontrino = []
    for row in db_dict['scontrino']:
        r_scontrino.append({'data': row['data'], 's1': row['s1'], 'd1': row['d1']})
        r_scontrino.append({'data': row['data'], 's1': row['s2'], 'd1': row['d2']})
        r_scontrino.append({'data': row['data'], 's1': row['s3'], 'd1': row['d3']})
    db_dict['scontrino'] = r_scontrino

    # Fill all the remaining tables
    for index, (key, value) in enumerate(db_dict.items(), start=1):
        table = document.tables[index]

        while len(table.rows) <= len(value):
            row = table.add_row()
            # row.height = Cm(0.61)

        for i, t in enumerate(value, start=1):
            table.cell(i, 0).text = t['data'].strftime('%d/%m/%Y')
            table.cell(i, 1).text = t['d1'] if t['d1'] is not None else ''
            table.cell(i, 2).text = t['s1'] if t['s1'] is not None else ''
            table.rows[i].height = Cm(0.61)

    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_parte_2_tmp.docx')
    output_name = f'Missione_{missione.id}_parte_2.docx'
    document.save(os.path.join(moduli_output_path, output_name_tmp))
    # Salvo il docx appena creato dentro a un FileField
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    outputStream = open(output_name_tmp, "rb")
    moduli_missione.parte_2_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def compila_autorizz_dottorandi(request, id):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, settings.STATIC_URL[1:], 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'AutorizzazioneDottorandi.pdf')
    coords_dict = {
        'data_richiesta': [120, 728],
        'tutor': [165, 577],
        'nomecognome': [410, 577],
        'anno_dottorato': [160, 550],
        'scuola_dottorato': [150, 522],
        'destinazione': [350, 438],
        'inizio': [100, 410],
        'fine': [250, 410],
        'motivazione': [62, 383],

    }
    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)

    value_dict = {
        'data_richiesta': date_richiesta.parte_1.strftime('%d/%m/%Y'),
        'tutor': profile.tutor,
        'nomecognome': f'{profile.user.first_name} {profile.user.last_name}',
        'anno_dottorato': f'{profile.anno_dottorato}',
        'scuola_dottorato': profile.scuola_dottorato,
        'destinazione': f'{missione.citta_destinazione} - {missione.stato_destinazione.nome}',
        'inizio': missione.inizio.strftime('%d/%m/%Y'),
        'fine': missione.fine.strftime('%d/%m/%Y'),
        'motivazione': missione.motivazione,
    }

    buffer = io.BytesIO()
    can = canvas.Canvas(buffer)
    can.setFont("Times-Roman", 11)
    # Write values
    for k, v in coords_dict.items():
        if k == 'cf':
            can.setFont('Courier', 18.2)
            can.drawString(*v, value_dict[k])
            can.setFont("Times-Roman", 11)
        else:
            if value_dict[k] is None:
                value_dict[k] = ""
            can.drawString(*v, value_dict[k])

    can.showPage()
    can.save()
    buffer.seek(0)
    new_pdf = PdfFileReader(buffer)

    # Leggo il file base
    input = PdfFileReader(open(input_file, "rb"))  # Base file
    page = input.getPage(0)
    # Faccio il merge delle modifiche con il file base
    page.mergePage(new_pdf.getPage(0))

    # Scrivo tutto in un file temporaneo
    output = PdfFileWriter()
    output.addPage(page)
    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_autoriz_dottorandi_tmp.pdf')
    outputStream = open(output_name_tmp, "wb")
    output.write(outputStream)
    outputStream.close()

    # Salvo il pdf appena creato dentro a un FileField
    output_name = f'Missione_{missione.id}_autoriz_dottorandi.pdf'
    outputStream = open(output_name_tmp, "rb")
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    moduli_missione.dottorandi_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)
